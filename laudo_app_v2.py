import streamlit as st
import re
from datetime import datetime
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import io
import shutil
import time
import traceback # Para detalhes de erro

# --- 1. Constantes (Com suas últimas modificações) ---
TIPOS_MATERIAL_BASE = {
    "v": "vegetal dessecado",
    "po": "pulverizado",
    "pd": "petrificado",
    "r": "resinoso"
}
TONALIDADES_ESPECIAIS = { # Não usado atualmente, mas mantido
    "b": ("branco", "esbranquiçada"),
    "a": ("amarelo", "amarelada")
}
TONALIDADES_GENERICAS = { # Usando suas abreviações
    "b": "esbranquiçada", "a": "amarelada", "vd": "esverdeada",
    "vr": "avermelhada", "az": "azulada", "p": "enegrecida",
    "c": "acinzentada", "m": "amarronzada", "r": "arrosada",
    "l": "alaranjada", "violeta": "arroxeadada"
}
TIPOS_EMBALAGEM_BASE = {
    "e": "microtubo do tipo “eppendorf”",
    "z": "embalagem do tipo \"zip\"",
    "a": "papel alumínio",
    "pl": "plástico",
    "pa": "papel"
}
CORES_FEMININO_EMBALAGEM = { # Inclui abreviações e nomes
    "t": "transparente",
    "branco": "branca", "branca": "branca", "b": "branca",
    "azul": "azul", "az": "azul",
    "amarelo": "amarela", "amarela": "amarela", "am": "amarela",
    "vermelho": "vermelha", "vermelha": "vermelha", "vm": "vermelha",
    "verde": "verde", "vd": "verde",
    "preto": "preta", "preta": "preta", "p": "preta", # Adicionado 'p'
    "cinza": "cinza", "c": "cinza",
    "marrom": "marrom", "m": "marrom",
    "rosa": "rosa", "r": "rosa",
    "laranja": "laranja", "l": "laranja",
    "violeta": "violeta",
    "roxa" : "roxa"
}
QUANTIDADES_EXTENSO = {
    1: "uma", 2: "duas", 3: "três", 4: "quatro", 5: "cinco",
    6: "seis", 7: "sete", 8: "oito", 9: "nove", 10: "dez",
    11: "onze", 12: "doze", 13: "treze", 14: "quatorze", 15: "quinze",
    16: "dezesseis", 17: "dezessete", 18: "dezoito", 19: "dezenove", 20: "vinte"
}
meses_portugues = {
    "January": "janeiro", "February": "fevereiro", "March": "março",
    "April": "abril", "May": "maio", "June": "junho", "July": "julho",
    "August": "agosto", "September": "setembro", "October": "outubro",
    "November": "novembro", "December": "dezembro"
}

# --- 2. Funções Auxiliares ---

def add_formatted_paragraph(document, text, font_name='Gadugi', font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, is_bold=False, is_italic=False, space_after=Pt(0)):
    """Adiciona um parágrafo formatado ao documento Word."""
    try:
        paragraph = document.add_paragraph()
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_after = space_after
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = is_bold
        run.italic = is_italic
        # Aplica fonte ao estilo do parágrafo para consistência
        # (Só se aplica bem se o parágrafo não tiver múltiplos runs com formatação diferente)
        # if len(paragraph.runs) == 1:
        #     try:
        #         style = paragraph.style
        #         font = style.font
        #         font.name = font_name
        #         font.size = Pt(font_size)
        #     except Exception as style_e:
        #         print(f"Debug: Erro ao aplicar estilo ao parágrafo: {style_e}") # Log para depuração
        # return paragraph
    except Exception as e:
        st.error(f"Erro em add_formatted_paragraph ao adicionar texto: '{text[:50]}...' - {e}")
        print(f"Erro em add_formatted_paragraph ao adicionar texto: '{text[:50]}...' - {e}") # Log console
        # Opcional: levantar o erro novamente se for crítico
        # raise e
    return paragraph # Retorna o parágrafo mesmo se o estilo falhar

def setup_default_font(document, font_name='Gadugi', font_size=12):
    """Define a fonte padrão para o documento."""
    try:
        style = document.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        # Tenta aplicar a outros estilos comuns
        for style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'Body Text']:
            if style_name in document.styles:
                style = document.styles[style_name]
                # Verifica se o estilo suporta formatação de fonte diretamente
                if hasattr(style, 'font'):
                     font = style.font
                     font.name = font_name
                     # Pode ajustar tamanhos específicos para cabeçalhos aqui se desejar
                     # if 'Heading' in style_name: font.size = Pt(14) else: font.size = Pt(font_size)
                     font.size = Pt(font_size) # Mantém 12pt para todos por enquanto
    except Exception as e:
        st.warning(f"Aviso: Não foi possível definir a fonte padrão para todos os estilos: {e}")
        print(f"Aviso: Não foi possível definir a fonte padrão para todos os estilos: {e}")


def pluralizar_palavra(palavra, quantidade):
    """Pluraliza uma palavra em português."""
    # Verifica se palavra é None ou não é string
    if not isinstance(palavra, str) or quantidade != 1:
        # Aplica regras de pluralização apenas se quantidade > 1
        if quantidade == 1 or not isinstance(palavra, str):
            return palavra # Retorna original se qtd=1 ou tipo inválido

        # Lógica de pluralização
        if palavra == "microtubo do tipo “eppendorf”": return palavra
        if palavra == "embalagem do tipo \"zip\"": return palavra
        if palavra == "microtubo do tipo “eppendorf”s": return "microtubo do tipo “eppendorf”"
        if palavra.endswith('m'): return re.sub(r'm$', 'ns', palavra)
        if palavra.endswith(('ão', 'ões')): return palavra # 'õe' é menos comum, talvez 'ões' cubra
        if palavra.endswith(('al', 'el', 'ol', 'ul')): return palavra[:-1] + 'is' # Regra do L -> is
        if palavra.endswith(('r', 'z', 's')): return palavra + 'es' # Regra R/Z/S -> es (cuidado com 's' já plural)
        # Simplificação: Adiciona 's' para a maioria das vogais e consoantes não cobertas acima
        # (Pode precisar de mais regras para casos específicos)
        return palavra + 's'
    return palavra # Retorna original se quantidade == 1

def obter_quantidade_extenso_web(quantidade):
    """Obtém a quantidade por extenso do dicionário."""
    if quantidade in QUANTIDADES_EXTENSO:
        return QUANTIDADES_EXTENSO[quantidade]
    else:
        print(f"AVISO (interno): Quantidade {quantidade} não mapeada para extenso. Usando numeral.")
        return str(quantidade)

def gerar_descricao_item_web(numero_item_str, item_data):
    """Gera a descrição formatada para o DOCX a partir dos dados do Streamlit."""
    try:
        quantidade = item_data.get('quantidade', 1) # Default 1 se faltar
        tipo_material_key = item_data.get('tipo_material')
        tipo_embalagem_base_key = item_data.get('tipo_embalagem_base')
        cor_embalagem_key_or_text = item_data.get('cor_embalagem')
        referencia_subitem = item_data.get('referencia_subitem', '[Ref. Ausente]') # Placeholder se faltar
        pessoa_relacionada = item_data.get('pessoa_relacionada')
        is_last_item = item_data.get('is_last', False)

        if not tipo_material_key or not tipo_embalagem_base_key:
            st.error(f"Erro interno: Tipo de material ou embalagem faltando para item {numero_item_str}")
            return f"[ERRO NA DESCRIÇÃO DO ITEM {numero_item_str}]"

        qtd_extenso = obter_quantidade_extenso_web(quantidade)
        porcoes = pluralizar_palavra("porção", quantidade)
        acondicionamento = "acondicionada em" if quantidade == 1 else "acondicionadas, individualmente, em"

        tipo_material_nome = TIPOS_MATERIAL_BASE.get(tipo_material_key, f"[Tipo Inválido: {tipo_material_key}]")
        embalagem_singular = TIPOS_EMBALAGEM_BASE.get(tipo_embalagem_base_key, f"[Emb. Inválida: {tipo_embalagem_base_key}]")

        # Monta a string da embalagem com cor, tratando None e strings vazias
        if tipo_embalagem_base_key in ["pl", "pa"] and cor_embalagem_key_or_text:
            # Busca pela chave OU usa o texto direto se não for chave conhecida
            cor_legivel = CORES_FEMININO_EMBALAGEM.get(cor_embalagem_key_or_text, cor_embalagem_key_or_text)
            # Tenta feminilizar se necessário (e se não for uma cor 'neutra' ou vazia)
            if cor_legivel and cor_legivel not in ['transparente', 'azul', 'verde', 'cinza', 'marrom', 'rosa', 'laranja', 'violeta', 'roxa'] and isinstance(cor_legivel, str) and cor_legivel.endswith('o'):
                cor_legivel = cor_legivel[:-1] + 'a'
            embalagem_singular = f"{embalagem_singular} de cor {cor_legivel}"

        # Pluraliza a embalagem (com ou sem cor)
        if "microtubo do tipo “eppendorf”" in embalagem_singular: emb_texto = embalagem_singular
        elif "embalagem do tipo \"zip\"" in embalagem_singular: emb_texto = embalagem_singular
        else:
            # Tenta aplicar pluralização mais cuidadosamente
            match_cor = re.search(r'^(.*?)( de cor .+)?$', embalagem_singular, re.IGNORECASE | re.UNICODE)
            if match_cor:
                 base = match_cor.group(1) # Parte antes da cor
                 cor_part = match_cor.group(2) if match_cor.group(2) else "" # Parte da cor
                 base_plural = pluralizar_palavra(base, quantidade)
                 emb_texto = f"{base_plural}{cor_part}"
            else: # Se não encontrar padrão com cor, pluraliza tudo
                 emb_texto = pluralizar_palavra(embalagem_singular, quantidade)


        referente = "referente" if quantidade == 1 else "referentes"
        terminacao = "." if is_last_item else ";"
        desc = f"{numero_item_str} {quantidade} ({qtd_extenso}) {porcoes} de material {tipo_material_nome}, {acondicionamento} {emb_texto}, {referente} à amostra do subitem {referencia_subitem} do laudo de constatação supracitado"
        desc += f", relacionada a {pessoa_relacionada}{terminacao}" if pessoa_relacionada else f"{terminacao}"
        return desc

    except Exception as e:
        st.error(f"Erro ao gerar descrição para item {numero_item_str}: {e}")
        print(f"Erro ao gerar descrição para item {numero_item_str}: {e}")
        return f"[ERRO NA DESCRIÇÃO DO ITEM {numero_item_str}]"

# --- Interface e Lógica Principal Streamlit ---

st.set_page_config(layout="wide")
st.title("Gerador de Laudo Pericial v2 (.docx)")

# Widget fora do formulário para definir o número de itens
# Usar um valor padrão e permitir atualização via widget
if 'num_itens' not in st.session_state:
    st.session_state['num_itens'] = 1 # Valor inicial padrão

num_itens = st.number_input(
    "Quantos itens deseja descrever?",
    min_value=1,
    value=st.session_state.num_itens, # Usa o valor do estado da sessão
    step=1,
    key='num_itens_selector',
    # Atualiza o estado da sessão quando o valor muda (Streamlit rerun)
    on_change=lambda: st.session_state.update(num_itens=st.session_state.num_itens_selector)
)


# Define as opções para os SelectBox
tipos_material_opcoes = {k: v for k, v in TIPOS_MATERIAL_BASE.items()}
tipos_embalagem_opcoes = {k: v for k, v in TIPOS_EMBALAGEM_BASE.items()}
# Usando CORES_FEMININO_EMBALAGEM para as opções de cor (chave: nome legível)
cores_opcoes_display = {k: v for k, v in CORES_FEMININO_EMBALAGEM.items()}
cores_opcoes_display["outra"] = "Outra (digitar)" # Adiciona opção extra


with st.form(key='laudo_form'):
    st.header("Custódia")
    lacre_num = st.text_input("Número do Lacre da Contraprova", placeholder="Ex: 0000659555", key='lacre')

    st.header("Informações dos Itens")
    itens_data = []
    # Loop baseado no num_itens definido ANTES do form
    for i in range(num_itens):
        st.subheader(f"Item 2.{i + 1}")
        # Usar colunas para melhor layout
        cols = st.columns([1, 3, 3]) # Proporção das colunas
        with cols[0]:
            qtd = st.number_input(f"Qtd. Porções", key=f'qtd_{i}', min_value=1, value=1, step=1)
        with cols[1]:
            tipo_mat = st.selectbox(f"Tipo Material", options=list(tipos_material_opcoes.keys()), format_func=lambda x: f"{x}: {tipos_material_opcoes[x]}", key=f'tipo_mat_{i}') # Mostra chave e nome
        with cols[2]:
            tipo_emb_key = st.selectbox(f"Tipo Embalagem", options=list(tipos_embalagem_opcoes.keys()), format_func=lambda x: f"{x}: {tipos_embalagem_opcoes[x]}", key=f'tipo_emb_{i}') # Mostra chave e nome

        cor_emb_final = None # Valor final da cor (chave ou texto)
        if tipo_emb_key in ['pl', 'pa']: # Só mostra opções de cor para plástico ou papel
            cols_cor = st.columns([1, 2]) # Colunas para selectbox e input de texto
            with cols_cor[0]:
                 # Usa as chaves do dict de cores como valor interno, mostra nome legível
                 cor_selecionada_key = st.selectbox(f"Cor Emb.", options=list(cores_opcoes_display.keys()), format_func=lambda x: cores_opcoes_display[x], key=f'cor_emb_{i}')
            if cor_selecionada_key == "outra":
                 with cols_cor[1]:
                    # Se for 'outra', pega o texto digitado
                    cor_emb_final = st.text_input("Digite a cor", key=f'cor_digitada_{i}').lower()
            else:
                 # Senão, guarda a chave da cor selecionada ('t', 'b', 'preta', etc)
                 cor_emb_final = cor_selecionada_key

        # Inputs restantes para o item
        ref_sub = st.text_input(f"Ref. Subitem Laudo Constatação", key=f'ref_{i}', placeholder="Ex: 2.1.1")
        pessoa_rel = st.text_input(f"Pessoa Relacionada (Opcional)", key=f'pessoa_{i}')

        # Guarda os dados coletados para este item
        itens_data.append({
            'quantidade': qtd,
            'tipo_material': tipo_mat,
            'tipo_embalagem_base': tipo_emb_key,
            'cor_embalagem': cor_emb_final,
            'referencia_subitem': ref_sub,
            'pessoa_relacionada': pessoa_rel,
            'is_last': (i == num_itens - 1)
        })

    st.header("Imagens (Opcional)")
    # Widget para upload de arquivos
    uploaded_files = st.file_uploader("Selecione as imagens (PNG, JPG, JPEG)",
                                       accept_multiple_files=True,
                                       type=['png', 'jpg', 'jpeg'],
                                       key='uploader')

    # Botão de submissão final do formulário
    submitted = st.form_submit_button("Gerar Laudo .docx")

# --- Lógica de Geração do DOCX (APÓS SUBMISSÃO) ---
if submitted:
    # Validação básica dos inputs obrigatórios
    valid = True
    if not lacre_num:
         st.error("Por favor, informe o número do Lacre.")
         valid = False
    if any(not item.get('referencia_subitem') for item in itens_data): # Checa se referencia existe e não é vazia
         st.error("Por favor, informe a Referência do Subitem para todos os itens.")
         valid = False

    if valid: # Prossegue somente se inputs básicos estão ok
        st.info("Gerando o documento Word... Aguarde.")
        document = Document()
        setup_default_font(document, font_name='Gadugi', font_size=12)
        temp_image_paths = [] # Lista para guardar caminhos de imgs temporárias
        temp_dir = "temp_images_laudo_streamlit_v2" # Pasta temporária

        try:
            # --- Início da Geração do Conteúdo DOCX ---

            # 2 MATERIAL RECEBIDO PARA EXAME
            add_formatted_paragraph(document, "2 MATERIAL RECEBIDO PARA EXAME", is_bold=True, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            tipos_material_itens_codigo = []
            subitens_cannabis = {}
            subitens_cocaina = {}
            has_cannabis_item = False
            has_cocaina_item = False

            # Adiciona a descrição de cada item ao DOCX
            for i, item_data in enumerate(itens_data):
                 numero_item_str = f"2.{i + 1}"
                 desc_item_txt = gerar_descricao_item_web(numero_item_str, item_data)
                 if "[ERRO" not in desc_item_txt: # Só adiciona se não deu erro na geração
                     add_formatted_paragraph(document, desc_item_txt, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
                 else:
                     st.error(f"Erro ao formatar item {numero_item_str}. Verifique os dados.")
                     # Pode parar a geração aqui se preferir: raise ValueError(...)

                 # Coleta dados para lógica das seções seguintes
                 tipo_cod = item_data['tipo_material']
                 ref_sub = item_data['referencia_subitem']
                 tipos_material_itens_codigo.append(tipo_cod)
                 if tipo_cod in ["v", "r"]:
                     if ref_sub and ref_sub not in subitens_cannabis: subitens_cannabis[ref_sub] = numero_item_str
                     has_cannabis_item = True
                 elif tipo_cod in ["po", "pd"]:
                     if ref_sub and ref_sub not in subitens_cocaina: subitens_cocaina[ref_sub] = numero_item_str
                     has_cocaina_item = True

            # Inserir Imagens carregadas
            imagens_inseridas_count = 0
            if uploaded_files:
                 st.write("Processando imagens carregadas...")
                 if not os.path.exists(temp_dir): os.makedirs(temp_dir)
                 for uploaded_file in uploaded_files:
                      # Usa um nome de arquivo temporário seguro (evita sobreposição)
                      base, ext = os.path.splitext(uploaded_file.name)
                      # Considerar usar uuid ou timestamp para nomes únicos se necessário
                      temp_path = os.path.join(temp_dir, f"{base}_{int(time.time()*1000)}{ext}")
                      with open(temp_path, "wb") as f: f.write(uploaded_file.getbuffer())
                      temp_image_paths.append(temp_path) # Guarda para limpar depois
                      if os.path.exists(temp_path):
                           try:
                                p = document.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = p.add_run()
                                run.add_picture(temp_path, width=Inches(5.5)) # Ajuste a largura se necessário
                                p.paragraph_format.space_after = Pt(6)
                                imagens_inseridas_count += 1
                           except Exception as e: st.warning(f"Erro ao inserir imagem '{uploaded_file.name}': {e}")
                      else: st.warning(f"Falha ao salvar temp '{uploaded_file.name}'.")

            # Adiciona Legenda da(s) Ilustração(ões)
            if imagens_inseridas_count > 0:
                caption_text = f"Ilustração 1 – Material recebido para exame." if imagens_inseridas_count == 1 else f"Ilustrações 1-{imagens_inseridas_count} – Material recebido para exame."
                add_formatted_paragraph(document, caption_text, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
            else: # Legenda padrão se não houver imagens
                add_formatted_paragraph(document, "Ilustração 1 – Material recebido para exame.", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

            # --- Seções 3 a 7, Referências, Data, Assinatura ---
            # 3 OBJETIVO DOS EXAMES
            add_formatted_paragraph(document, "3 OBJETIVO DOS EXAMES", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            add_formatted_paragraph(document, "Visa esclarecer à autoridade requisitante quanto às características do material apresentado, bem como se ele contém substância de uso proscrito no Brasil e capaz de causar dependência física e/ou psíquica. O presente laudo pericial busca demonstrar a materialidade da infração penal apurada.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(12))

            # 4 EXAMES
            add_formatted_paragraph(document, "4 EXAMES", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            has_cannabis_code = any(item in ["v", "r"] for item in tipos_material_itens_codigo)
            has_cocaina_code = any(item in ["po", "pd"] for item in tipos_material_itens_codigo)
            section_index_cocaina = 4.1
            if has_cannabis_code:
                section_index_cocaina = 4.2
                add_formatted_paragraph(document, "4.1 Exames realizados para pesquisa de Cannabis Sativa L. (maconha)", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
                add_formatted_paragraph(document, "4.1.1 Ensaio químico com Fast blue salt B: teste de cor em reação com solução aquosa de sal de azul sólido B em meio alcalino;", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(0))
                add_formatted_paragraph(document, "4.1.2 Cromatografia em Camada Delgada (CCD), comparativa com substância padrão, em sistemas contendo eluentes apropriados e posterior revelação com solução aquosa de azul sólido B.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            if has_cocaina_code:
                section_prefix = f"{section_index_cocaina:.1f}"
                add_formatted_paragraph(document, f"{section_prefix} Exames realizados para pesquisa de cocaína", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
                add_formatted_paragraph(document, f"{section_prefix}.1 Ensaio químico com teste de tiocianato de cobalto-reação de cor com solução de tiocianato de cobalto em meio ácido;", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(0))
                add_formatted_paragraph(document, f"{section_prefix}.2 Cromatografia em Camada Delgada (CCD), comparativa com substância padrão, em sistemas com eluentes apropriados e revelação com solução de iodo platinado.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            elif not has_cannabis_code and not has_cocaina_code:
                add_formatted_paragraph(document, "4.1 Exames realizados", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
                add_formatted_paragraph(document, "4.1.1 Exame macroscópico;", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            if document.paragraphs: document.paragraphs[-1].paragraph_format.space_after = Pt(12)

            # 5 RESULTADOS
            add_formatted_paragraph(document, "5 RESULTADOS", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            result_section_index_cocaina = 5.1
            if has_cannabis_item:
                result_section_index_cocaina = 5.2
                subitens_ref_cannabis_list = sorted(subitens_cannabis.keys())
                subitens_texto_cannabis = " e ".join(subitens_ref_cannabis_list) if len(subitens_ref_cannabis_list) > 1 else "".join(subitens_ref_cannabis_list)
                ref_label = "subitem" if len(subitens_ref_cannabis_list) == 1 else "subitens"
                add_formatted_paragraph(document, f"5.1 Resultados obtidos para os materiais descritos no(s) {ref_label} {subitens_texto_cannabis}:", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
                add_formatted_paragraph(document, "5.1.1 No ensaio com Fast blue salt B, foram obtidas coloração característica para canabinol e tetrahidrocanabinol (princípios ativos da Cannabis sativa L.).", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(0))
                add_formatted_paragraph(document, "5.1.2 Na CCD, obtiveram-se perfis cromatográficos coincidentes com o material de referência (padrão de Cannabis sativa L.); portanto, a substância tetrahidrocanabinol está presente nos materiais questionados.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            if has_cocaina_item:
                section_prefix = f"{result_section_index_cocaina:.1f}"
                subitens_ref_cocaina_list = sorted(subitens_cocaina.keys())
                subitens_texto_cocaina = " e ".join(subitens_ref_cocaina_list) if len(subitens_ref_cocaina_list) > 1 else "".join(subitens_ref_cocaina_list)
                ref_label = "subitem" if len(subitens_ref_cocaina_list) == 1 else "subitens"
                add_formatted_paragraph(document, f"{section_prefix} Resultados obtidos para os materiais descritos no(s) {ref_label} {subitens_texto_cocaina}:", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
                add_formatted_paragraph(document, f"{section_prefix}.1 No teste de tiocianato de cobalto, foram obtidas coloração característica para cocaína;", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(0))
                add_formatted_paragraph(document, f"{section_prefix}.2 Na CCD, obteve-se perfis cromatográficos coincidentes com o material de referência (padrão de cocaína); portanto, a substância cocaína está presente nos materiais questionados.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            if document.paragraphs: document.paragraphs[-1].paragraph_format.space_after = Pt(12)

            # 6 CONCLUSÃO
            add_formatted_paragraph(document, "6 CONCLUSÃO", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            conclusoes = []
            if has_cannabis_item:
                subitens_ref_cannabis_list = sorted(subitens_cannabis.keys())
                subitens_texto_cannabis = " e ".join(subitens_ref_cannabis_list) if len(subitens_ref_cannabis_list) > 1 else "".join(subitens_ref_cannabis_list)
                ref_label = "subitem" if len(subitens_ref_cannabis_list) == 1 else "subitens"
                conclusion_cannabis = f"nos materiais descritos no(s) {ref_label} {subitens_texto_cannabis}, foi detectada a presença de partes da planta Cannabis sativa L., vulgarmente conhecida por maconha. A Cannabis sativa L. contém princípios ativos chamados canabinóis, dentre os quais se encontra o tetrahidrocanabinol, substância perturbadora do sistema nervoso central. Tanto a Cannabis sativa L. quanto a tetrahidrocanabinol são proscritas no país, com fulcro na Portaria nº 344/1998, atualizada por meio da RDC nº 970, de 19/03/2025, da Anvisa."
                conclusoes.append(conclusion_cannabis)
            if has_cocaina_item:
                subitens_ref_cocaina_list = sorted(subitens_cocaina.keys())
                subitens_texto_cocaina = " e ".join(subitens_ref_cocaina_list) if len(subitens_ref_cocaina_list) > 1 else "".join(subitens_ref_cocaina_list)
                ref_label = "subitem" if len(subitens_ref_cocaina_list) == 1 else "subitens"
                conclusion_cocaina = f"nos materiais descritos no(s) {ref_label} {subitens_texto_cocaina}, foi detectada a presença de cocaína, substância alcaloide estimulante do sistema nervoso central. A cocaína é proscrita no país, com fulcro na Portaria nº 344/1998, atualizada por meio da RDC nº 970, de 19/03/2025, da Anvisa."
                conclusoes.append(conclusion_cocaina)
            if len(conclusoes) == 2: texto_final_conclusao = f"A partir das análises realizadas, conclui-se que, {conclusoes[0]} Outrossim, {conclusoes[1]}"
            elif len(conclusoes) == 1: texto_final_conclusao = f"A partir das análises realizadas, conclui-se que, {conclusoes[0]}"
            else: texto_final_conclusao = "A partir das análises realizadas, conclui-se que não foram detectadas substâncias de uso proscrito nos materiais analisados, com fulcro na Portaria nº 344/1998, atualizada por meio da RDC nº 970, de 19/03/2025, da Anvisa."
            add_formatted_paragraph(document, texto_final_conclusao, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(12))

            # 7 CUSTÓDIA DO MATERIAL
            add_formatted_paragraph(document, "7 CUSTÓDIA DO MATERIAL", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            add_formatted_paragraph(document, "7.1 Contraprova", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            texto_lacre_docx = f"7.1.1 A amostra contraprova ficará armazenada neste Instituto, conforme Portaria 0003/2019/SSP (Lacre nº {lacre_num})." # Usa lacre_num do form
            add_formatted_paragraph(document, texto_lacre_docx, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(12))

            # REFERÊNCIAS
            add_formatted_paragraph(document, "REFERÊNCIAS", is_bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            referencias_base = [
                "BRASIL. Ministério da Saúde. Portaria SVS/MS n° 344, de 12 de maio de 1998. Aprova o regulamento técnico sobre substâncias e medicamentos sujeitos a controle especial. Diário Oficial da União: Brasília, DF, p. 37, 19 maio 1998. Alterada pela RDC nº 970, de 19/03/2025.",
                "GOIÁS. Secretaria de Estado da Segurança Pública. Portaria nº 0003/2019/SSP de 10 de janeiro de 2019. Regulamenta a apreensão, movimentação, exames, acondicionamento, armazenamento e destruição de drogas no âmbito da Secretaria de Estado da Segurança Pública. Diário Oficial do Estado de Goiás: n° 22.972, Goiânia, GO, p. 4-5, 15 jan. 2019.",
                "SWGDRUG: Scientific Working Group for the Analysis of Seized Drugs. Recommendations. Version 8.0 june. 2019. Disponível em: http://www.swgdrug.org/Documents/SWGDRUG%20Recommendations%20Version%208_FINAL_ForPosting_092919.pdf. Acesso em: 07/10/2019."
            ]
            if has_cannabis_item: referencias_base.append("UNODC (United Nations Office on Drugs and Crime). Laboratory and scientific section. Recommended Methods for the Identification and Analysis of Cannabis and Cannabis Products. New York: 2012.")
            if has_cocaina_item:
                 ref_cocaina = "UNODC (United Nations Office on Drugs and Crime). Laboratory and Scientific Section. Recommended Methods for the Identification and Analysis of Cocaine in Seized Materials. New York: 2012."
                 # Evita adicionar duplicado se ambos existirem
                 if ref_cocaina not in [ref for ref in referencias_base if 'Cocaine' in ref]:
                      referencias_base.append(ref_cocaina)
            for ref in referencias_base:
                add_formatted_paragraph(document, ref, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6))
            if document.paragraphs: document.paragraphs[-1].paragraph_format.space_after = Pt(12)

            # Encerramento
            add_formatted_paragraph(document, "É o que se tem a relatar.", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(24))

            # --- Data ---
            hoje = datetime.now()
            data_formatada = f"Goiânia, {hoje.day} de {meses_portugues.get(hoje.strftime('%B'), hoje.strftime('%B'))} de {hoje.year}."
            add_formatted_paragraph(document, data_formatada, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(24))

            # --- Assinatura ---
            add_formatted_paragraph(document, "Laudo assinado digitalmente com dados do assinador à esquerda das páginas", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(0))
            add_formatted_paragraph(document, "Daniel Chendes Lima", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0)) # Centralizado
            add_formatted_paragraph(document, "Perito Criminal", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12)) # Centralizado


            # --- Salvar em memória para download ---
            bio = io.BytesIO()
            document.save(bio)
            bio.seek(0)

            st.success("Laudo gerado com sucesso!")
            st.download_button(
                label="Baixar Laudo (.docx)",
                data=bio,
                file_name="laudo_pericial_v2.docx", # Nome do arquivo atualizado
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key='download_button'
            )

        except Exception as e:
            st.error(f"Ocorreu um erro durante a geração do documento: {e}")
            st.error(traceback.format_exc()) # Mostra mais detalhes do erro

        finally:
            # --- Limpeza dos arquivos temporários de imagem ---
            if temp_image_paths:
                st.write("Limpando imagens temporárias...")
                if os.path.exists(temp_dir):
                     try:
                          shutil.rmtree(temp_dir) # Remove a pasta e todo o conteúdo
                          st.write("Limpeza concluída.")
                     except Exception as e:
                          st.warning(f"Erro ao limpar pasta temporária '{temp_dir}': {e}")
